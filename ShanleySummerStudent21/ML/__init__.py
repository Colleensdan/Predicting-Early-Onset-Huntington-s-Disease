import os
import numpy as np
import pandas as pd
import glob
import matplotlib.pyplot as plt
from os import chdir
import sklearn
from sklearn import metrics, datasets, neighbors
from sklearn.datasets import make_classification
from sklearn.ensemble import RandomForestClassifier
from sklearn.inspection import permutation_importance
import sys
import warnings
from sklearn.model_selection import train_test_split
import itertools
from sklearn.metrics import plot_confusion_matrix
from sklearn.metrics import accuracy_score
from sklearn.metrics import balanced_accuracy_score
from sklearn.metrics import cohen_kappa_score
from sklearn.metrics import classification_report
from sklearn.metrics import roc_curve
from sklearn.metrics import average_precision_score
from sklearn.metrics import RocCurveDisplay
from sklearn.metrics import precision_recall_curve
from sklearn.metrics import PrecisionRecallDisplay
import numpy as np
import matplotlib.pyplot as plt
from sklearn.metrics import accuracy_score
from sklearn.metrics import balanced_accuracy_score
from sklearn.metrics import cohen_kappa_score
from sklearn.metrics import classification_report
from sklearn.metrics import roc_curve
from sklearn.metrics import average_precision_score
from sklearn.metrics import RocCurveDisplay
from sklearn.metrics import precision_recall_curve
from sklearn.metrics import PrecisionRecallDisplay
import numpy as np
import matplotlib.pyplot as plt
from sklearn.naive_bayes import GaussianNB
from sklearn import svm, datasets
from sklearn.model_selection import train_test_split
from sklearn.metrics import plot_confusion_matrix
from sklearn.model_selection import learning_curve
from sklearn.svm import SVC
from sklearn import svm, datasets
from sklearn.model_selection import train_test_split
from sklearn.metrics import plot_confusion_matrix
from sklearn.model_selection import cross_val_score

from sklearn.model_selection import learning_curve
from sklearn.svm import SVC
from sklearn import svm
from ML import *
from docx import Document
from docx.shared import Inches
import matplotlib.pyplot as plt
import numpy as np
import io
from ML import plot_decision_boundaries

warnings.filterwarnings("ignore")
np.set_printoptions(precision=2)

# Global variables are always in caps, to distinguish them from local variables.
WORKING_DIRECTORY = os.path.dirname(__file__)
DATA_DIR = os.path.join(WORKING_DIRECTORY, 'Data')
INFO_DIR = os.path.join(WORKING_DIRECTORY, 'Info')
PLOT_DIR = os.path.join(INFO_DIR, "PLOT")
Data = os.path.join(DATA_DIR, "ML_Data.csv")
ValData = os.path.join(DATA_DIR, "ML_Data_val.csv")

# Flags that change the behaviour of the control_script

# flag to demonstrate the principle of flags
Import = True
ImportVal = True
DataInvestigation = False
split = True
featureSelection = False
printSelected = False
printNotSelected = False
scaler = True
Alg_tests = True
plotAlg = True
crossVal = True
confusionMatrix = True
RobustML = False



def get_X_y(df):
    """
    Separates the phenotypes from the data, and tidies unnecessary columns
    :param df (DataFrame) containing both targets and data with features
    :return: X (DataFrame) tidied data with no targets
    :return: y (DataFrame) targets of X (either HD or WT)
    """
    X = df.loc[:, df.columns != 'Conditions']
    X = X.loc[:, X.columns != "Unnamed: 0"]
    y = df.loc[:, df.columns == 'Conditions']
    return X, y

def get_age_files(x_f, filename, remove_duplicates=True):
    """
    Splits a dataframe into the X and y components and into test and training sets

    Attributes:
        :param x_f: file to read
        :param filename: str - filename of file to read
        :param remove_duplicates: (bool) if true, duplicates are removed. Keep to reduce overfitting

        :return: X_train (DataFrame)
        :return X_test (DataFrame)
        :return y_train (DataFrame)
        :return y_test (DataFrame)
    """
    X_train = pd.read_csv(x_f)
    y_n = filename.replace("X", "y")
    y_train = pd.read_csv(y_n)
    val = pd.read_csv((filename.replace("X_", "")).replace("train", "validation"))
    X_test, y_test = get_X_y(val)
    X_test = X_test.drop(columns="Samples")


    if remove_duplicates:
        X_train=X_train.drop(columns="Unnamed: 0").drop_duplicates()
        y_train = X_train.join(y_train)
        y_train=y_train["Conditions"]
        y_train= pd.DataFrame(y_train)

    else:
        y_train=y_train.drop(columns="Unnamed: 0")
        X_train = X_train.drop(columns="Unnamed: 0")
    val_counts=y_train.value_counts()

    # find rows where there is only one of HD or WT
    if 1 in val_counts.values:
        condition = val_counts.isin([1])

        # find the unique category
        _ = str(condition[condition==True].index.values.tolist())
        # df containing nans is produced except for the phenotype (HD/WT)
        cat = _.replace("[('", "").replace("',)]", "")
        index = (y_train[y_train == cat]).dropna().index.values.astype(int)
        # duplicate the category for y
        row = pd.Series({"Conditions": cat})
        y_train = y_train.append(row, ignore_index=True)
        # duplicate the category for x
        row = X_train.iloc[index]
        X_train=X_train.append(row, ignore_index=True)

    return X_train, X_test, y_train, y_test


def permutation_based_feature_importance(clf, X, y, feature_names, X_t, y_t, save=False, filename = None, loc = ""):
    """
    Produces plots of feature importance based on prior classification
    :param clf: Obj - a trained classifier
    :param X: DataFrame - test samples against features
    :param y: DataFrame - test targets
    :param feature_names: list - names of the biomarkers
    :param X_t: DataFrame - training samples
    :param y_t: DataFrame - training targets
    :param save: bool - True is parsing into .docx, if False will output figure
    :param filename: str - name to save file, only needed if save=True
    :param loc: str - location for file to be saved, only needed if save=True
    """

    #clf = GaussianNB()
    clf.fit(X_t, y_t)

    result = permutation_importance(clf, X, y)

    fig, ax = plt.subplots()
    sorted_idx = result.importances_mean.argsort()
    ax.boxplot(result.importances[sorted_idx].T,
               vert=False, labels=feature_names[sorted_idx])
    ax.set_title("Permutation Importance of each feature")
    ax.set_ylabel("Features")
    fig.tight_layout()
    if not save:
        plt.show()
    else:
        if not filename:
            raise FileNotFoundError ("Provide a filename to save the file")
        else:
            plt.savefig(filename)

    if loc != "":
        raise NameError("Implement loc here")

def plot_confusion(classifier, X_test, y_test, class_names):
    """
    Plots a confusion matrix
    :param classifier: (Object) trained classifier
    :param X_test: DataFrame containing samples
    :param y_test: DataFrame containing targets
    :param class_names: list, either HD or WT
    :return:
    """
    # https://scikit-learn.org/stable/auto_examples/model_selection/plot_confusion_matrix.html#sphx-glr-auto-examples-model-selection-plot-confusion-matrix-py

    np.set_printoptions(precision=2)

    # Plot non-normalized confusion matrix
    titles_options = [("Confusion matrix, without normalization", None),
                      ("Normalized confusion matrix", 'true')]
    for title, normalize in titles_options:
        disp = plot_confusion_matrix(classifier, X_test, y_test,
                                     display_labels=class_names,
                                     cmap=plt.cm.Blues,
                                     normalize=normalize)
        disp.ax_.set_title(title)

        print(title)
        print(disp.confusion_matrix)

    plt.show()


def evaluate_model(y_pred, y_true, X_test, y_test, clf, target_names, X_train, y_train, print_scores = False, document=None, fname=None):
    """
    Evaluates the success of a model, using various techniques including numerical evaluations looking at accuracy,
    F1, Fbeta among others, as well as confusion matrices and permutation importances

    :param y_pred: (DataFrame) predicted targets based off of previous classification
    :param y_true: (DataFrame) actual targets from previous classification
    :param X_test: (DataFrame) test data containing samples
    :param y_test: (DataFrame) test targets
    :param clf: (Object) a trained classifier
    :param target_names: (list) either HD or WT
    :param X_train: (DataFrame) containing samples
    :param y_train: (DataFrame) containing targets
    :param print_scores: (bool) if true, then files won't be saved, and no .docx will be produced
    :param document: (Object) .docx file to be produced
    :param fname: (str) filename to be saved
    """
    if print_scores:
        ######################################################
        # accuracy
        print("Accuracy: ", accuracy_score(y_true, y_pred))
        ###################################################
        # balanced accuracy
        print("Balanced accuracy score: ", balanced_accuracy_score(y_true, y_pred))
        #########################
        # cohen_kappa_score
        """
        The kappa score is a number between -1 and 1. Scores above .8 are generally considered good agreement; zero or lower means no agreement (practically random labels)
        """
        print("cohen kappa score: ",cohen_kappa_score(y_true, y_pred), "above 0.8 is good agreement")
        ##############################
        # plot confusion matrix
        plot_confusion(clf, X_test, y_test, ["HD", "WT"])
        ####################################
        # classification report

        print("classification report: \n", classification_report(y_true, y_pred, target_names=target_names))
        #########################################
        # general metrics
        print("Precision: ",metrics.precision_score(y_true, y_pred, average="binary", pos_label="HD"))

        print("Recall:", metrics.recall_score(y_true, y_pred, average="binary", pos_label="HD"))

        print("F1:",metrics.f1_score(y_true, y_pred, average="binary", pos_label="HD"))

        print("F beta, beta-0.5", metrics.fbeta_score(y_true, y_pred, beta=0.5,average="binary", pos_label="HD"))

        print("F beta, beta-1",metrics.fbeta_score(y_true, y_pred, beta=1,average="binary", pos_label="HD"))

        print("F beta, beta-2",metrics.fbeta_score(y_true, y_pred, beta=2,average="binary", pos_label="HD"))

        print("precision recall fscore support", metrics.precision_recall_fscore_support(y_true, y_pred, beta=0.5,average="binary", pos_label="HD"))


        # ROC curve
        y_scores = clf.predict_proba(X_test)[:, 1]
        precision, recall, threshold = precision_recall_curve(y_true, y_scores, pos_label="HD")


        print("Average precision score: ", average_precision_score(y_true, y_scores, pos_label="HD"))

    if document is not None:
        if fname is None:
            raise NameError("Provide a filename to save this document")
        document.add_heading("Test Metrics", level=2)
        document.add_paragraph(("Accuracy: {}".format(accuracy_score(y_true, y_pred))), style = "List Bullet")
        document.add_paragraph(("Balanced accuracy score: {}".format(balanced_accuracy_score(y_true, y_pred))), style = "List Bullet")
        document.add_paragraph(("Cohen kappa score: {} ".format(accuracy_score(y_true, y_pred))), style = "List Bullet")
        p=document.add_paragraph("", style = "List Bullet")
        p.add_run('(The kappa score is a number between -1 and 1. Scores above .8 are generally considered good agreement; zero or lower means no agreement (practically random labels)).').italic = True


        # confusion matricies
        document.add_heading("Confusion Matrices", level=2)

        np.set_printoptions(precision=2)

        # Plot confusion matrices
        titles_options = [("Confusion matrix, without normalization", None),
                          ("Normalized confusion matrix", 'true')]
        for title, normalize in titles_options:
            memfile = io.BytesIO()
            disp = plot_confusion_matrix(clf, X_test, y_test,
                                         display_labels=["HD", "WT"],
                                         cmap=plt.cm.Blues,
                                         normalize=normalize)
            disp.ax_.set_title(title)

            plt.savefig(memfile)
            document.add_picture(memfile, width=Inches(5))
            memfile.close()

        # classification report
        document.add_heading("Classification report", level=2)
        document.add_paragraph("{}".format(classification_report(y_true, y_pred, target_names=target_names)))

        # Precision/recall
        document.add_heading("Precision/Recall Scores", level=2)
        document.add_paragraph("Precision: {}".format(metrics.precision_score(y_true, y_pred, average="binary", pos_label="HD")), style= "List Bullet")
        document.add_paragraph("Recall: {}".format(metrics.recall_score(y_true, y_pred, average="binary", pos_label="HD")), style= "List Bullet")
        document.add_paragraph("F1 {}".format(metrics.f1_score(y_true, y_pred, average="binary", pos_label="HD")), style= "List Bullet")

        # Decision boundaries plot

        document.add_heading("Decision Surface of model - training")
        memfile = io.BytesIO()
        m = clf
        pca_clf = plot_decision_boundaries.DecisionBoundaries(model=m, name=fname).plot(X_train, y_train, memfile)
        plt.savefig(memfile)
        document.add_picture(memfile, width=Inches(5))
        memfile.close()

        """
        # todo - Krutik, I can't imagine I will have time to finish this section. If you want to create figures to show the test data on the decision surface, i think you need to adjust the dimensions of the plot within plot_decision_boundaries.DecisionBoundaries(), so they are the same as on the first plot, thus, the decision surface will be comparable for both plots 
        
        document.add_heading("Decision Surface of model - testing")
        memfile2 = io.BytesIO()
        plot_decision_boundaries.DecisionBoundaries(model=pca_clf, name=fname).test_plot(pca_clf, X_test, y_test, memfile2, X_train, y_train)
        plt.savefig(memfile2)
        document.add_picture(memfile2, width=Inches(5))
        memfile2.close()

        """

        # feature importance

        memfile = io.BytesIO()
        y = permutation_based_feature_importance(clf, X_test, y_test, X_train.columns, X_train, y_train, save=True, filename = memfile)
        document.add_picture(memfile, width=Inches(5))
        memfile.close()

        document.save(r'../../ML/Classifiers/{}.docx'.format(fname))
        print("Saved {}.docx".format(fname), "in ../../ML/Classifiers/")
