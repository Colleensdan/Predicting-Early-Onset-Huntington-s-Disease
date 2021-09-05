from sklearn.naive_bayes import GaussianNB
from ML import *
from docx import Document

# todo both svm_classifier.py and naive_bayes_classifier.py should inherit from an ABC classifier

def classifyNB(f, filename, save_document=True):
    """

    Train a gaussian naive bayes classifier

    Take care when returning variables, since if save_doc=true the document object will be returned, otherwise there
    will be one less variable returned

    Attributes:
    :param f: (Object - TextIOWrapper) Data file to be opened
    :param filename: (str) Filename of datafile
    :param save_document: (bool) Will save the outputs to a docx file if true

    :return opt (Object) an optimised SVC
    :return X_test (DataFrame) Test data with appropriate biomarkers
    :return y_test (DataFrame) The phenotypes [HD/WT] of the samples associated with X_test
        :return y_preds (Dataframe) The predicted phenotypes of the SVC
    :return ["HD", "WT"] (list) The possible phenotypes
    :return document (Object) The .docx file to be created. Only returned on save_document=True
    :return rna (DataFrame) an aggregate of all data (training and test, with conditions tied to samples in a new column)
    :return X_train (DataFrame) Training data with appropriate biomarkers
    :return y_train (DataFrame) The predicted phenotypes of X_train
"""
    clf = GaussianNB()

    X_train, X_test, y_train, y_test = get_age_files(f, filename, remove_duplicates=False)

    X_train, X_test = __removeknownfeatures(X_train, X_test, filename)
    clf.fit(X_train, y_train)

    y_preds = clf.predict(X_test)

    rna = filename.replace("X_", "").replace(".csv", "")

    if save_document:
        document = Document()
        document.add_heading("Gaussian Naive Baiyes Classifier Performance on {} data".format(filename.replace("X_", "").replace("_", " ").replace("train.csv", ""), "classifier"),0)
        document.add_heading("Training", level=2)
        document.add_paragraph(('Training Accuracy : %.3f'%clf.score(X_train, y_train)))
        return clf, X_test, y_test, y_preds, ["HD", "WT"], document, rna, X_train, y_train

    print(filename,'\nTraining Accuracy : %.3f'%clf.score(X_train, y_train))
    return clf, X_test, y_test, y_preds, ["HD", "WT"], rna,X_train, y_train


def __removeknownfeatures(X_train, X_test, fname):
    """
Private method which removes miRNAs which have been known to be unhelpful in classification due to a negative, or zero permutation score

Attributes:
    :param X_train (DataFrame) training data without the associated conditions
    :param X_test (DataFrame) test data without the associated conditions
    :param fname (str) current filename e.g. X_miRNA_10m_train

    :return X_train (DataFrame) X_train with  appropriate biomarkers
    :return X_test (DataFrame) X_test with appropriate biomarkers
"""
    if "miRNA" in fname:
        if "10m"  in fname:
            print(fname)
            X_train = X_train.drop(columns=["mmu-miR-7037-3p",
                                            "mmu-miR-3103-5p",
                                            "mmu-miR-6998-5p",
                                            "mmu-miR-693-3p"])


            X_test = X_test.drop(columns=["mmu-miR-7037-3p",
                                          "mmu-miR-3103-5p",
                                          "mmu-miR-6998-5p",
                                          "mmu-miR-693-3p"])
        if "2m" in fname:
            X_train = X_train.drop(columns=["mmu-miR-465a-3p",
                                        "mmu-miR-465c-3p",
                                        "mmu-miR-465b-3p",
                                        "mmu-miR-20b-3p",
                                        "mmu-miR-7008-5p",
                                        "mmu-miR-1982-5p",
                                        "mmu-miR-12182-3p",
                                        "mmu-miR-10a-3p"])
            X_test = X_test.drop(columns=["mmu-miR-465a-3p",
                                      "mmu-miR-465c-3p",
                                      "mmu-miR-465b-3p",
                                      "mmu-miR-20b-3p",
                                      "mmu-miR-7008-5p",
                                      "mmu-miR-1982-5p",
                                      "mmu-miR-12182-3p",
                                      "mmu-miR-10a-3p"])
    return X_train, X_test


"""
# Example usage if running in isolation

for filename in glob.glob('*X*'):
    with open(os.path.join(os.getcwd(), filename), 'r') as f:
        clf, X_test, y_test, y_preds, classes ,document, name,X_train, y_train = classifyNB(f, filename, save_document=True)
        evaluate_model(y_preds, y_test, X_test, y_test, clf, classes, X_train, y_train, document=document, fname=name.replace("_train", ""))
"""
