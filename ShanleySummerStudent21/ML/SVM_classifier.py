from skopt import BayesSearchCV
from ML import *
from docx import Document

from sklearn.linear_model import LogisticRegression
from sklearn.model_selection import train_test_split, GridSearchCV, StratifiedKFold
from sklearn.preprocessing import MinMaxScaler
from imblearn.over_sampling import SMOTE
from imblearn.pipeline import Pipeline as imbpipeline
from RScripts import improve_group_imbalances



def bayes_search_CV_init(self, estimator, search_spaces, optimizer_kwargs=None,
                         n_iter=50, scoring=None, fit_params=None, n_jobs=1,
                         n_points=1, iid=True, refit=True, cv=None, verbose=0,
                         pre_dispatch='2*n_jobs', random_state=None,
                         error_score='raise', return_train_score=False):
    """
    A hacky fix such that bayes_search can be used due to some deprecation.
    """

    self.search_spaces = search_spaces
    self.n_iter = n_iter
    self.n_points = n_points
    self.random_state = random_state
    self.optimizer_kwargs = optimizer_kwargs
    self._check_search_space(self.search_spaces)
    self.fit_params = fit_params

    super(BayesSearchCV, self).__init__(
        estimator=estimator, scoring=scoring,
        n_jobs=n_jobs, refit=refit, cv=cv, verbose=verbose,
        pre_dispatch=pre_dispatch, error_score=error_score,
        return_train_score=return_train_score)

BayesSearchCV.__init__ = bayes_search_CV_init

def classifySVM(f, filename, save_document=True):
    """
    Classify an optimised SVM using Bayes Search with cross validaton (in order to find the optimal hyper-parameters for this learner)

    Take care when returning variables, since if save_doc is true the document object will be returned, otherwise there
    will be one less variable returned

    Attributes:
        :param f: (Object - TextIOWrapper) Data file to be opened
        :param filename: (str) Filename of datafile
        :param save_document: (bool) Will save the outputs to a docx file if true

        :return opt (Object) an optimised SVC
        :return X_test (DataFrame) Test data with appropriate biomarkers
        :return y_test (DataFrame) The phenotypes [HD/WT] of the samples associated with X_test
        :return y_preds (Dataframe) The predicted phenotypes of the SVC
        :return ["HD", "WT"] (list) The possible phenotypes
        :return document (Object) The .docx file to be created. Only returned on save_document=True
        :return rna (DataFrame) an aggregate of all data (training and test, with conditions tied to samples in a new column)
        :return X_train (DataFrame) Training data with appropriate biomarkers
        :return y_train (DataFrame) The predicted phenotypes of X_train
    """
    X_train, X_test, y_train, y_test = get_age_files(f, filename, remove_duplicates=False)

    X_train, X_test = __removeknownfeatures(X_train, X_test, filename)
    # log-uniform search: search over p = exp(x) by varying x

    # There are some issues with optimising SVM using BayesSearchCV
    # https://github.com/scikit-optimize/scikit-optimize/issues/1006
    # ensure you are running 0.23.2
    # https://scikit-learn.org/stable/install.html

    opt = BayesSearchCV(
        svm.SVC(probability=True),
        {
            'C': (1e-6, 1e+6, 'log-uniform'),
            'gamma': (1e-6, 1e+1, 'log-uniform'),
            'degree': (1, 8),  # integer valued parameter
            'kernel': ['linear', 'poly', 'rbf'],
            # categorical parameter
        },
        n_iter=32,
        cv=3
    )

    opt.fit(X_train, y_train)
    """
    # Hi Krutik, I had some issues implementing SMOTE within a pipeline. I thought I would leave in a partial 
    solution in case you have better luck/more time than me. I'm not sure how important this is, since this analysis 
    uses accuracy scores and not so much CV scores. Best of luck, and apologies for not being able to complete this part.
    
    
    pipeline = imbpipeline(steps = [['smote', improve_group_imbalances.SMOTE(save=False, X=X_train, y=y_train).perform_smote(X_train, y_train)],
                                    ['scaler', MinMaxScaler()],
                                    ['classifier', LogisticRegression(random_state=11,
                                                                      max_iter=1000)]])

    stratified_kfold = StratifiedKFold(n_splits=3,
                                       shuffle=True,
                                       random_state=11)

    param_grid = {'classifier__C':[0.001, 0.01, 0.1, 1, 10, 100, 1000]}
    grid_search = GridSearchCV(estimator=pipeline,
                               param_grid=param_grid,
                               scoring='roc_auc',
                               cv=stratified_kfold,
                               n_jobs=-1)

    grid_search.fit(X_train, y_train)
    cv_score = grid_search.best_score_
    test_score = grid_search.score(X_test, y_test)
    print(f'Cross-validation score: {cv_score}\nTest score: {test_score}')
"""
    # use this line if using an optimised SVC
    print("val. score: %s" % opt.best_score_)


    print("test score: %s" % opt.score(X_test, y_test))
    y_preds = opt.predict(X_test)

    rna = filename.replace("X_", "").replace(".csv", "")

    if save_document:
        document = Document()
        document.add_heading("Classifier Performance on {} data".format(filename.replace("X_", "").replace("_", " ").replace("train.csv", ""), "classifier"),0)
        document.add_heading("Training", level=2)
        document.add_paragraph(('Training Accuracy : %.3f'%opt.score(X_train, y_train)))
        return opt, X_test, y_test, y_preds, ["HD", "WT"], document, rna, X_train, y_train

    print(filename,'\nTraining Accuracy : %.3f'%opt.score(X_train, y_train))
    return opt, X_test, y_test, y_preds, ["HD", "WT"], rna,X_train, y_train



def __removeknownfeatures(X_train, X_test, fname):
    """
    Private method which removes miRNAs which have been known to be unhelpful in classification due to a negative, or zero permutation score

    Attributes:
        :param X_train (DataFrame) training data without the associated conditions
        :param X_test (DataFrame) test data without the associated conditions
        :param fname (str) current filename e.g. X_miRNA_10m_train

        :return X_train (DataFrame) X_train with  appropriate biomarkers
        :return X_test (DataFrame) X_test with appropriate biomarkers
    """
    if "miRNA" in fname:
        if "2m" in fname:
            X_train = X_train.drop(columns=["mmu-miR-7008-5p",
                                            "mmu-miR-20b-3p"])

            X_test = X_test.drop(columns=["mmu-miR-7008-5p",
                                          "mmu-miR-20b-3p"])

            return X_train, X_test

        if "10m" in fname:
            X_train = X_train.drop(columns=["mmu-miR-7037-3p",
                                            "mmu-miR-3103-5p",
                                            "mmu-miR-6998-5p",
                                            "mmu-miR-693-3p"])

            X_test = X_test.drop(columns=["mmu-miR-7037-3p",
                                          "mmu-miR-3103-5p",
                                          "mmu-miR-6998-5p",
                                          "mmu-miR-693-3p"])

            return X_train, X_test

    return X_train, X_test